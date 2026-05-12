"""
Document Download Service
--------------------------
Converts optimized markdown content back to a .docx file.
Preserves professional formatting and SEO structure.
"""

import re
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def markdown_to_docx(markdown_content: str, filename: str = "seo_optimized") -> bytes:
    """
    Convert optimized markdown content to a formatted .docx file.
    Returns bytes ready for download.
    """
    doc = Document()
    _set_document_styles(doc)

    lines = markdown_content.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip empty lines (handled by paragraph spacing)
        if not line.strip():
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            para = doc.add_heading(line[2:].strip(), level=1)
            _style_heading(para, level=1)

        elif line.startswith("## "):
            para = doc.add_heading(line[3:].strip(), level=2)
            _style_heading(para, level=2)

        elif line.startswith("### "):
            para = doc.add_heading(line[4:].strip(), level=3)
            _style_heading(para, level=3)

        # Bullet lists
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")

        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line).strip()
            doc.add_paragraph(text, style="List Number")

        # Horizontal rule
        elif line.strip() == "---":
            _add_horizontal_line(doc)

        # Placeholder comments (shown in a distinct style)
        elif line.strip().startswith("[") and line.strip().endswith("]"):
            para = doc.add_paragraph(line.strip())
            for run in para.runs:
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # green
                run.font.italic = True

        # Table (basic markdown table)
        elif "|" in line and i + 1 < len(lines) and "|" in lines[i + 1] and "---" in lines[i + 1]:
            i = _add_table(doc, lines, i)
            continue

        # Regular paragraph
        else:
            clean = _clean_inline_markdown(line.strip())
            if clean:
                para = doc.add_paragraph()
                _add_formatted_run(para, clean)

        i += 1

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _set_document_styles(doc: Document):
    """Set document-wide styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)


def _style_heading(para, level: int):
    """Apply custom heading colors."""
    colors = {
        1: RGBColor(0x1A, 0x1A, 0x2E),  # dark navy
        2: RGBColor(0x16, 0x21, 0x3E),  # navy
        3: RGBColor(0x0F, 0x3D, 0x6B),  # blue
    }
    for run in para.runs:
        run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))


def _add_horizontal_line(doc: Document):
    """Add a horizontal rule paragraph."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_table(doc: Document, lines: list, start_i: int) -> int:
    """Parse and add a markdown table. Returns the new line index."""
    # Find table rows
    table_lines = []
    i = start_i
    while i < len(lines) and "|" in lines[i]:
        if "---" not in lines[i]:
            table_lines.append(lines[i])
        i += 1

    if len(table_lines) < 1:
        return i

    # Parse cells
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return i

    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"

    for r_idx, row_cells in enumerate(rows):
        for c_idx, cell_text in enumerate(row_cells):
            if c_idx < max_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = _clean_inline_markdown(cell_text)
                # Bold header row
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

    doc.add_paragraph()  # spacing after table
    return i


def _clean_inline_markdown(text: str) -> str:
    """Remove inline markdown syntax (bold, italic, code) for plain text."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # bold
    text = re.sub(r'\*(.+?)\*', r'\1', text)        # italic
    text = re.sub(r'`(.+?)`', r'\1', text)          # code
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # links → anchor text
    return text


def _add_formatted_run(para, text: str):
    """Add a run with basic bold/italic formatting from markdown."""
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
        elif part:
            para.add_run(part)
